import logging
import os
import uuid
from typing import Any

import chromadb
from chromadb.config import Settings
from fastapi import HTTPException, UploadFile
from openai import OpenAI
from PyPDF2 import PdfReader
from sqlmodel import Session, select

from app.core.config import settings
from app.models import Document, DocumentChunk, User
from app.utils import get_file_hash

logger = logging.getLogger(__name__)


class DocumentService:
    def __init__(self, db: Session):
        self.db = db
        self.upload_dir = settings.UPLOAD_DIR
        os.makedirs(self.upload_dir, exist_ok=True)
        self.client = OpenAI(api_key=settings.OPENAI_API_KEY)

        # Initialize Chroma client (optional - fallback if not available)
        self.chroma_client = None
        self.collection = None
        try:
            self.chroma_client = chromadb.HttpClient(
                host="chroma",
                port=8000,
                settings=Settings(
                    allow_reset=True,
                    anonymized_telemetry=False
                )
            )

            # Create or get collection
            self.collection = self.chroma_client.get_or_create_collection(
                name="documents",
                metadata={"hnsw:space": "cosine"}
            )
        except Exception as e:
            logger.warning(f"ChromaDB not available, using database-only mode: {e}")
            self.chroma_client = None
            self.collection = None

    async def upload_document(
        self, file: UploadFile, user: User, description: str | None = None
    ) -> Document:
        """Upload a document and create initial database record."""
        if not file.content_type == "application/pdf":
            raise HTTPException(status_code=400, detail="Only PDF files are supported")

        # Read file content and validate
        content = await file.read()
        if len(content) > settings.MAX_UPLOAD_SIZE:
            raise HTTPException(status_code=400, detail="File too large")

        # Generate unique filename using UUID and original extension
        file_hash = get_file_hash(content)
        ext = os.path.splitext(file.filename)[1]
        filename = f"{file_hash}{ext}"
        file_path = os.path.join(self.upload_dir, filename)

        # Save file to disk
        with open(file_path, "wb") as f:
            f.write(content)

        # Create document record
        document = Document(
            filename=filename,
            original_filename=file.filename,
            file_size=len(content),
            content_type=file.content_type,
            description=description,
            file_path=file_path,
            user_id=user.id,
        )
        self.db.add(document)
        self.db.commit()
        self.db.refresh(document)

        # Start async processing
        await self.process_document(document)

        return document

    async def process_document(self, document: Document) -> None:
        """Process uploaded PDF document and create chunks with embeddings."""
        try:
            # Update status to processing
            document.processing_status = "processing"
            self.db.commit()

            # Extract text from PDF
            chunks = self._extract_text_chunks(document.file_path)

            # Process chunks and store in Chroma (if available)
            for idx, (content, metadata) in enumerate(chunks):
                # Generate embedding using OpenAI (if ChromaDB is available)
                embedding = None
                if self.collection:
                    try:
                        embedding = await self.generate_embedding(content)

                        # Store in Chroma
                        self.collection.add(
                            documents=[content],
                            embeddings=[embedding],
                            metadatas=[{
                                "document_id": str(document.id),
                                "user_id": str(document.user_id),
                                "chunk_index": idx,
                                **metadata
                            }],
                            ids=[f"{document.id}_{idx}"]
                        )
                    except Exception as e:
                        logger.warning(f"Failed to store in ChromaDB: {e}")

                # Store in database
                chunk = DocumentChunk(
                    document_id=document.id,
                    user_id=document.user_id,
                    content=content,
                    chunk_index=idx,
                    chunk_metadata=metadata,
                )
                self.db.add(chunk)

            # Update document status and chunk count
            document.processing_status = "completed"
            document.chunk_count = len(chunks)
            self.db.commit()

        except Exception as e:
            document.processing_status = "failed"
            self.db.commit()
            raise HTTPException(
                status_code=500,
                detail=f"Failed to process document: {str(e)}",
            )

    async def generate_embedding(self, text: str) -> list[float]:
        """Generate embedding for text using OpenAI."""
        try:
            response = self.client.embeddings.create(
                model="text-embedding-ada-002",
                input=text
            )
            return response.data[0].embedding
        except Exception as e:
            logger.error(f"Error generating embedding: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail="Failed to generate embedding"
            )

    def _extract_text_chunks(self, file_path: str) -> list[tuple[str, str]]:
        """Extract text from PDF and split into chunks with metadata."""
        chunks = []

        try:
            with open(file_path, "rb") as file:
                pdf_reader = PdfReader(file)

                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()

                    # Split text into chunks (simple implementation - can be enhanced)
                    words = text.split()
                    chunk_size = 200  # Adjust based on your needs

                    for i in range(0, len(words), chunk_size):
                        chunk_words = words[i:i + chunk_size]
                        chunk_text = " ".join(chunk_words)

                        metadata = {
                            "page": page_num + 1,
                            "chunk_start": i,
                            "chunk_end": i + len(chunk_words),
                        }

                        chunks.append((chunk_text, str(metadata)))

        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to extract text from PDF: {str(e)}",
            )

        return chunks

    def get_document(self, document_id: uuid.UUID, user_id: uuid.UUID) -> Document:
        """Get a document by ID and user ID."""
        document = self.db.exec(
            select(Document).where(
                Document.id == document_id,
                Document.user_id == user_id,
            )
        ).first()

        if not document:
            raise HTTPException(status_code=404, detail="Document not found")

        return document

    def get_user_documents(
        self, user_id: uuid.UUID, skip: int = 0, limit: int = 100
    ) -> tuple[list[Document], int]:
        """Get all documents for a user with pagination."""
        query = select(Document).where(Document.user_id == user_id)
        total = len(self.db.exec(query).all())

        documents = self.db.exec(
            query.offset(skip).limit(limit)
        ).all()

        return documents, total

    def delete_document(self, document_id: uuid.UUID, user_id: uuid.UUID) -> None:
        """Delete a document and its associated file."""
        document = self.get_document(document_id, user_id)

        # Delete file from disk
        try:
            if os.path.exists(document.file_path):
                os.remove(document.file_path)
        except Exception as e:
            # Log error but continue with database deletion
            print(f"Error deleting file {document.file_path}: {e}")

        # Delete from Chroma
        try:
            self.collection.delete(
                where={"document_id": str(document_id)}
            )
        except Exception as e:
            # Log error but continue with database deletion
            print(f"Error deleting vectors from Chroma: {e}")

        # Delete from database (will cascade to chunks)
        self.db.delete(document)
        self.db.commit()

    async def search_documents(
        self,
        session: Session,
        user_id: str,
        query: str,
        limit: int = 5
    ) -> list[dict[str, Any]]:
        """Search documents using semantic similarity."""
        try:
            # Generate query embedding
            query_embedding = await self.generate_embedding(query)

            # Search in Chroma
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=limit,
                where={"user_id": user_id}
            )

            # Format results
            formatted_results = []
            for idx, (doc, distance) in enumerate(zip(results["documents"][0], results["distances"][0], strict=False)):
                metadata = results["metadatas"][0][idx]
                formatted_results.append({
                    "content": doc,
                    "similarity": 1 - distance,  # Convert distance to similarity
                    "metadata": metadata,
                    "chunk": await self._get_chunk(
                        session,
                        metadata["document_id"],
                        metadata["chunk_index"]
                    )
                })

            return formatted_results

        except Exception as e:
            logger.error(f"Error searching documents: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error searching documents: {str(e)}"
            )

    async def _get_chunk(
        self,
        session: Session,
        document_id: str,
        chunk_index: int
    ) -> DocumentChunk:
        """Get a specific document chunk."""
        chunk = session.exec(
            select(DocumentChunk)
            .where(
                DocumentChunk.document_id == uuid.UUID(document_id),
                DocumentChunk.chunk_index == chunk_index
            )
        ).first()

        if not chunk:
            raise HTTPException(
                status_code=404,
                detail=f"Chunk not found: doc={document_id}, idx={chunk_index}"
            )

        return chunk

    def extract_text_from_file(self, file_path: str) -> str:
        """
        Synchronous version of text extraction for Celery tasks.
        
        Args:
            file_path: Path to the file to extract text from
            
        Returns:
            Extracted text content
        """
        try:
            if file_path.lower().endswith('.pdf'):
                return self._extract_pdf_text_sync(file_path)
            elif file_path.lower().endswith(('.doc', '.docx')):
                return self._extract_docx_text_sync(file_path)
            elif file_path.lower().endswith('.txt'):
                return self._extract_txt_text_sync(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_path}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise

    def _extract_pdf_text_sync(self, file_path: str) -> str:
        """Synchronous PDF text extraction."""
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(file_path)
            text = ""

            for page in doc:
                text += page.get_text()

            doc.close()
            return text

        except ImportError:
            # Fallback to PyPDF2
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)
            text = ""

            for page in reader.pages:
                text += page.extract_text()

            return text

    def _extract_docx_text_sync(self, file_path: str) -> str:
        """Synchronous DOCX text extraction."""
        try:
            from docx import Document

            doc = Document(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            return text
        except ImportError:
            logger.warning("python-docx not available - cannot extract DOCX text")
            return ""

    def _extract_txt_text_sync(self, file_path: str) -> str:
        """Synchronous text file extraction."""
        with open(file_path, encoding='utf-8') as file:
            return file.read()
